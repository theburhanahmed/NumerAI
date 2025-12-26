"""
Enhanced reports service for custom reports, templates, scheduling, and exports.
"""
from typing import Dict, List, Any, Optional
from datetime import datetime, timedelta
from django.utils import timezone
from reports.models import ReportTemplate, GeneratedReport, ScheduledReport, ReportComparison
from numerology.models import Person, PersonNumerologyProfile
from reports.report_generator import generate_report_content
import json


class ReportsService:
    """Service for enhanced report generation and management."""
    
    def generate_custom_report(
        self,
        person: Person,
        numerology_profile: PersonNumerologyProfile,
        sections: List[str],
        custom_config: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Generate a custom report with selected sections.
        
        Args:
            person: Person instance
            numerology_profile: PersonNumerologyProfile instance
            sections: List of section names to include
            custom_config: Optional custom configuration (colors, fonts, etc.)
            
        Returns:
            Dictionary with custom report content
        """
        # Default sections if none provided
        if not sections:
            sections = ['numbers', 'interpretations', 'summary']
        
        content = {
            'person_name': person.name,
            'birth_date': person.birth_date.isoformat(),
            'report_type': 'custom',
            'generated_at': timezone.now().isoformat(),
            'sections': {},
            'custom_config': custom_config or {}
        }
        
        # Generate numbers section
        if 'numbers' in sections:
            content['sections']['numbers'] = {
                'life_path': numerology_profile.life_path_number,
                'destiny': numerology_profile.destiny_number,
                'soul_urge': numerology_profile.soul_urge_number,
                'personality': numerology_profile.personality_number,
                'attitude': numerology_profile.attitude_number,
                'maturity': numerology_profile.maturity_number,
                'balance': numerology_profile.balance_number,
                'personal_year': numerology_profile.personal_year_number,
                'personal_month': numerology_profile.personal_month_number,
            }
        
        # Generate interpretations section
        if 'interpretations' in sections:
            from numerology.interpretations import get_interpretation
            interpretations = {}
            number_fields = ['life_path', 'destiny', 'soul_urge', 'personality']
            
            for field in number_fields:
                number_value = getattr(numerology_profile, f'{field}_number')
                try:
                    interpretations[field] = get_interpretation(number_value)
                except ValueError:
                    interpretations[field] = None
            
            content['sections']['interpretations'] = interpretations
        
        # Generate summary section
        if 'summary' in sections:
            content['sections']['summary'] = f"Custom numerology report for {person.name} generated on {timezone.now().strftime('%Y-%m-%d')}"
        
        # Add other sections based on request
        available_sections = {
            'charts': self._generate_charts_section,
            'cycles': self._generate_cycles_section,
            'remedies': self._generate_remedies_section,
        }
        
        for section in sections:
            if section in available_sections:
                content['sections'][section] = available_sections[section](person, numerology_profile)
        
        return content
    
    def create_report_template(
        self,
        user,
        name: str,
        description: str,
        report_type: str,
        template_config: Dict[str, Any],
        is_premium: bool = False
    ) -> ReportTemplate:
        """
        Create a custom report template.
        
        Args:
            user: User instance
            name: Template name
            description: Template description
            report_type: Report type
            template_config: Template configuration
            is_premium: Whether template is premium
            
        Returns:
            Created ReportTemplate instance
        """
        template = ReportTemplate.objects.create(
            name=name,
            description=description,
            report_type=report_type,
            is_custom=True,
            owner=user,
            template_config=template_config,
            is_premium=is_premium,
            is_active=True
        )
        
        return template
    
    def schedule_report(
        self,
        user,
        template: ReportTemplate,
        person: Person,
        schedule_frequency: str,
        next_run_date: datetime
    ) -> ScheduledReport:
        """
        Schedule a recurring report.
        
        Args:
            user: User instance
            template: ReportTemplate instance
            person: Person instance
            schedule_frequency: Frequency (daily, weekly, monthly, yearly, custom)
            next_run_date: When to run next
            
        Returns:
            Created ScheduledReport instance
        """
        scheduled = ScheduledReport.objects.create(
            user=user,
            template=template,
            person=person,
            schedule_frequency=schedule_frequency,
            next_run_date=next_run_date,
            is_active=True
        )
        
        return scheduled
    
    def compare_reports(
        self,
        user,
        report1: GeneratedReport,
        report2: GeneratedReport
    ) -> Dict[str, Any]:
        """
        Compare two reports side-by-side.
        
        Args:
            user: User instance
            report1: First GeneratedReport instance
            report2: Second GeneratedReport instance
            
        Returns:
            Dictionary with comparison analysis
        """
        # Extract numbers from both reports
        content1 = report1.content if isinstance(report1.content, dict) else json.loads(report1.content) if isinstance(report1.content, str) else {}
        content2 = report2.content if isinstance(report2.content, dict) else json.loads(report2.content) if isinstance(report2.content, str) else {}
        
        numbers1 = content1.get('numbers', {})
        numbers2 = content2.get('numbers', {})
        
        # Compare numbers
        number_comparison = {}
        differences = []
        similarities = []
        
        all_keys = set(numbers1.keys()) | set(numbers2.keys())
        
        for key in all_keys:
            val1 = numbers1.get(key)
            val2 = numbers2.get(key)
            
            if val1 == val2:
                similarities.append(key)
                number_comparison[key] = {
                    'report1': val1,
                    'report2': val2,
                    'match': True
                }
            else:
                differences.append(key)
                number_comparison[key] = {
                    'report1': val1,
                    'report2': val2,
                    'match': False,
                    'difference': abs(val1 - val2) if val1 and val2 else None
                }
        
        # Calculate similarity score
        similarity_score = (len(similarities) / len(all_keys) * 100) if all_keys else 0
        
        comparison_data = {
            'report1_id': str(report1.id),
            'report1_title': report1.title,
            'report2_id': str(report2.id),
            'report2_title': report2.title,
            'number_comparison': number_comparison,
            'similarities': similarities,
            'differences': differences,
            'similarity_score': round(similarity_score, 2),
            'compared_at': timezone.now().isoformat()
        }
        
        # Save comparison
        comparison = ReportComparison.objects.create(
            user=user,
            report1=report1,
            report2=report2,
            comparison_data=comparison_data
        )
        
        return comparison_data
    
    def export_report_multiple_formats(
        self,
        report: GeneratedReport,
        format_type: str
    ) -> bytes:
        """
        Export report in multiple formats.
        
        Args:
            report: GeneratedReport instance
            format_type: Format (pdf, docx, json, html)
            
        Returns:
            Bytes content of exported report
        """
        if format_type == 'json':
            import json
            return json.dumps(report.content, indent=2).encode('utf-8')
        
        elif format_type == 'html':
            return self._export_as_html(report)
        
        elif format_type == 'docx':
            return self._export_as_docx(report)
        
        elif format_type == 'pdf':
            # Use existing PDF export
            from reports.views import export_generated_report_pdf
            # This would need to be refactored to return bytes
            # For now, return None to indicate use existing endpoint
            return None
        
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _generate_charts_section(
        self,
        person: Person,
        numerology_profile: PersonNumerologyProfile
    ) -> Dict[str, Any]:
        """Generate charts section data."""
        return {
            'chart_types': ['bar', 'pie'],
            'data': {
                'numbers': {
                    'life_path': numerology_profile.life_path_number,
                    'destiny': numerology_profile.destiny_number,
                    'soul_urge': numerology_profile.soul_urge_number,
                }
            }
        }
    
    def _generate_cycles_section(
        self,
        person: Person,
        numerology_profile: PersonNumerologyProfile
    ) -> Dict[str, Any]:
        """Generate cycles section data."""
        return {
            'personal_year': numerology_profile.personal_year_number,
            'personal_month': numerology_profile.personal_month_number,
            'cycles': []
        }
    
    def _generate_remedies_section(
        self,
        person: Person,
        numerology_profile: PersonNumerologyProfile
    ) -> Dict[str, Any]:
        """Generate remedies section data."""
        return {
            'remedies': [],
            'suggestions': []
        }
    
    def _export_as_html(self, report: GeneratedReport) -> bytes:
        """Export report as HTML."""
        content = report.content if isinstance(report.content, dict) else json.loads(report.content) if isinstance(report.content, str) else {}
        
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <title>{report.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; }}
        .section {{ margin: 20px 0; }}
        .number {{ font-size: 24px; font-weight: bold; color: #4a5568; }}
    </style>
</head>
<body>
    <h1>{report.title}</h1>
    <p>Generated for: {report.person.name}</p>
    <p>Date: {report.generated_at.strftime('%Y-%m-%d')}</p>
    <div class="content">
        {self._format_content_as_html(content)}
    </div>
</body>
</html>
"""
        return html.encode('utf-8')
    
    def _format_content_as_html(self, content: Dict[str, Any]) -> str:
        """Format content dictionary as HTML."""
        html_parts = []
        
        if 'sections' in content:
            for section_name, section_data in content['sections'].items():
                html_parts.append(f'<div class="section"><h2>{section_name.title()}</h2>')
                if isinstance(section_data, dict):
                    for key, value in section_data.items():
                        if isinstance(value, (int, float)):
                            html_parts.append(f'<div class="number">{key}: {value}</div>')
                        else:
                            html_parts.append(f'<p>{key}: {value}</p>')
                html_parts.append('</div>')
        
        return ''.join(html_parts)
    
    def _export_as_docx(self, report: GeneratedReport) -> bytes:
        """Export report as DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            doc.add_heading(report.title, 0)
            doc.add_paragraph(f"Generated for: {report.person.name}")
            doc.add_paragraph(f"Date: {report.generated_at.strftime('%Y-%m-%d')}")
            
            content = report.content if isinstance(report.content, dict) else json.loads(report.content) if isinstance(report.content, str) else {}
            
            if 'sections' in content:
                for section_name, section_data in content['sections'].items():
                    doc.add_heading(section_name.title(), level=1)
                    if isinstance(section_data, dict):
                        for key, value in section_data.items():
                            doc.add_paragraph(f"{key}: {value}", style='List Bullet')
            
            # Save to BytesIO
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
        
        except ImportError:
            raise ImportError("python-docx library is required for DOCX export")
